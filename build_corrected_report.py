from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    ListFlowable,
    ListItem,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "report_output"
OUT_DIR.mkdir(exist_ok=True)

PDF_PATH = OUT_DIR / "Отчет_Weather_Forecast_App_исправленный.pdf"
DOCX_PATH = OUT_DIR / "Отчет_Weather_Forecast_App_исправленный.docx"

FONT = "ArialReport"
FONT_BOLD = "ArialReportBold"
pdfmetrics.registerFont(TTFont(FONT, r"C:\Windows\Fonts\arial.ttf"))
pdfmetrics.registerFont(TTFont(FONT_BOLD, r"C:\Windows\Fonts\arialbd.ttf"))


SECTIONS = [
    "ВВЕДЕНИЕ",
    "1. ОБЩАЯ ХАРАКТЕРИСТИКА ПРОЕКТА",
    "2. АНАЛИЗ СТРУКТУРЫ ИСХОДНОГО КОДА",
    "3. РАЗРАБОТКА СЕРВЕРНОЙ ЧАСТИ",
    "4. РАЗРАБОТКА КЛИЕНТСКОЙ ЧАСТИ",
    "5. ОФОРМЛЕНИЕ И АДАПТИВНЫЙ ИНТЕРФЕЙС",
    "6. ТЕСТИРОВАНИЕ И КОНТРОЛЬ КАЧЕСТВА",
    "7. ИТОГИ И ВЫВОДЫ",
    "8. ПРИЛОЖЕНИЕ",
]

GOALS = [
    "реализовать серверное приложение FastAPI и подключить шаблон Jinja2;",
    "создать endpoint поиска городов через Open-Meteo Geocoding API;",
    "создать endpoint получения текущей погоды, почасового прогноза и прогноза на 10 дней;",
    "разработать мобильный HTML-интерфейс с экранами главной страницы, поиска, настроек и ошибки сети;",
    "реализовать выбор города, единиц измерения, языка, темы и города по умолчанию;",
    "добавить русский и английский язык, включая поиск русских названий городов;",
    "реализовать погодные фоны, которые меняются по фактическому состоянию погоды и времени суток;",
    "проверить работоспособность API, статических файлов, JavaScript и пользовательских сценариев.",
]

FEATURE_ROWS = [
    ["Функция", "Реализация в проекте"],
    ["Текущая погода", "Температура, влажность, скорость ветра, состояние погоды, максимум и минимум дня."],
    ["Почасовой прогноз", "Backend отдаёт 24 ближайших часа, frontend показывает горизонтально прокручиваемую ленту."],
    ["Прогноз на 10 дней", "Ежедневный список с погодной иконкой, min/max температурой, осадками и ветром."],
    ["Поиск города", "Маршрут /api/search выполняет геокодирование через Open-Meteo."],
    ["Русский/английский язык", "Параметр language передаётся на backend; поиск «Москва» работает на русском."],
    ["Единицы измерения", "Переключение °C/°F меняет temperature_unit и wind_speed_unit в запросах."],
    ["Город по умолчанию", "Default City выбирается в настройках и сохраняется в localStorage."],
    ["Погодные фоны", "Фоны из static/assets выбираются по WMO-коду погоды и времени суток."],
    ["Offline-сценарий", "При сетевой ошибке отображается экран Network error с кнопкой Retry."],
]

FILE_ROWS = [
    ["Файл/каталог", "Назначение"],
    ["main.py", "Точка входа FastAPI; маршруты /, /api/search, /api/weather; Open-Meteo; WMO-коды; ошибки."],
    ["templates/index.html", "Единый HTML-шаблон: главная, поиск, настройки, ошибка сети, нижняя навигация."],
    ["static/app.js", "Запросы API, localStorage, локализация, темы, выбор города, погодные фоны, DOM-отрисовка."],
    ["static/styles.css", "Premium minimal UI, glassmorphism, адаптивность, светлая/тёмная тема."],
    ["static/assets/", "Локальные погодные фоны: ясный день, ясная ночь, облачность, дождь, дождь ночью."],
    ["requirements.txt", "fastapi, uvicorn, jinja2, python-multipart, httpx."],
    ["README.md", "Описание возможностей, структуры и запуска проекта."],
]

TECH_ROWS = [
    ["Компонент", "Роль"],
    ["Python / FastAPI", "Серверная часть, маршрутизация и JSON API."],
    ["Uvicorn", "ASGI-сервер для локального запуска."],
    ["Jinja2", "Возврат HTML-страницы приложения."],
    ["httpx", "Асинхронные HTTP-запросы к Open-Meteo."],
    ["HTML5 / CSS3", "Структура экранов и визуальная система."],
    ["JavaScript", "Клиентские сценарии, запросы, localStorage, DOM-отрисовка."],
    ["Open-Meteo APIs", "Геокодирование городов и погодный прогноз."],
]

API_ROWS = [
    ["Endpoint", "Параметры", "Назначение"],
    ["GET /", "-", "Возвращает HTML-страницу приложения через Jinja2Templates."],
    ["GET /api/search", "query, limit, language", "Ищет города и возвращает название, страну, регион, координаты и часовой пояс."],
    ["GET /api/weather", "city, units, language", "Возвращает город, текущую погоду, 24 часа прогноза, прогноз на 10 дней и единицы измерения."],
]

TEST_ROWS = [
    ["Проверка", "Результат"],
    ["Импорт приложения", "main.app успешно импортируется, приложение называется Weather Forecast App."],
    ["Python compile", "python -m compileall выполняется без синтаксических ошибок."],
    ["JavaScript syntax", "static/app.js проходит node --check."],
    ["Главная страница", "GET / возвращает HTTP 200."],
    ["Статические файлы", "styles.css, app.js и изображения из static/assets возвращают HTTP 200."],
    ["Получение погоды", "/api/weather?city=Bryansk&units=celsius возвращает прогноз."],
    ["Русский поиск", "/api/search?query=Москва&language=ru возвращает город Москва."],
    ["Русский прогноз", "/api/weather?city=Москва&language=ru возвращает русское название и 24 часа прогноза."],
    ["Смена единиц", "При units=fahrenheit возвращаются °F и mph."],
]


def esc(text: str) -> str:
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")


def build_pdf() -> None:
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("Body", fontName=FONT, fontSize=10.5, leading=15, alignment=TA_JUSTIFY, spaceAfter=7))
    styles.add(ParagraphStyle("Small", fontName=FONT, fontSize=8.8, leading=12, alignment=TA_LEFT, spaceAfter=4))
    styles.add(ParagraphStyle("TitleRu", fontName=FONT_BOLD, fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=14))
    styles.add(ParagraphStyle("H1Ru", fontName=FONT_BOLD, fontSize=14, leading=18, textColor=colors.HexColor("#1F4D78"), spaceBefore=10, spaceAfter=8))
    styles.add(ParagraphStyle("H2Ru", fontName=FONT_BOLD, fontSize=12, leading=16, textColor=colors.HexColor("#2E74B5"), spaceBefore=8, spaceAfter=6))
    styles.add(ParagraphStyle("Center", fontName=FONT, fontSize=10.5, leading=15, alignment=TA_CENTER, spaceAfter=7))
    styles.add(ParagraphStyle("CodeRu", fontName=FONT, fontSize=8.5, leading=11, leftIndent=8, backColor=colors.HexColor("#F4F6F9"), borderColor=colors.HexColor("#DADCE0"), borderWidth=0.4, borderPadding=6, spaceAfter=8))

    def p(text: str, style: str = "Body") -> Paragraph:
        return Paragraph(esc(text), styles[style])

    def bullets(items):
        return ListFlowable([ListItem(p(item)) for item in items], bulletType="bullet", leftIndent=18)

    def table(rows, widths=None):
        data = [[p(cell, "Small") for cell in row] for row in rows]
        widths = widths or ([5.0 * cm, 11.4 * cm] if len(rows[0]) == 2 else [3.8 * cm, 4.4 * cm, 8.2 * cm])
        t = Table(data, colWidths=widths, hAlign="LEFT", repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CAD3DF")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return t

    class Report(BaseDocTemplate):
        pass

    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setFont(FONT, 8.5)
        canvas.setFillColor(colors.HexColor("#555555"))
        canvas.drawString(2 * cm, A4[1] - 1.25 * cm, "ОП-02068025.09.02.07.085.25")
        canvas.drawRightString(A4[0] - 2 * cm, A4[1] - 1.25 * cm, f"Лист {doc.page}")
        canvas.setStrokeColor(colors.HexColor("#DADCE0"))
        canvas.line(2 * cm, A4[1] - 1.42 * cm, A4[0] - 2 * cm, A4[1] - 1.42 * cm)
        canvas.restoreState()

    doc = Report(str(PDF_PATH), pagesize=A4, rightMargin=2 * cm, leftMargin=2 * cm, topMargin=2 * cm, bottomMargin=1.8 * cm)
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="normal")
    doc.addPageTemplates([PageTemplate(id="main", frames=frame, onPage=on_page)])

    story = []
    story += [
        p("Министерство науки и высшего образования Российской Федерации", "Center"),
        p("ФГБОУ ВО «Брянский государственный инженерно-технологический университет»", "Center"),
        p("Многопрофильный колледж", "Center"),
        p("Кафедра «Информационные технологии»", "Center"),
        Spacer(1, 1 * cm),
        p("ОТЧЕТ", "TitleRu"),
        p("по учебной практике по технологии разработки программного обеспечения", "Center"),
        p("Тема проекта: «Weather Forecast App - веб-приложение прогноза погоды»", "Center"),
        p(r"Локальный проект: C:\Users\siefi\Documents\New project\weather_forecast_app", "Center"),
        Spacer(1, 0.8 * cm),
    ]
    for line in ["Студент ______________________________", "Группа ______________________________", "№ зачётной книжки ______________________________", "Руководитель от вуза ______________________________", "Нормоконтроль ______________________________"]:
        story.append(p(line))
    story += [Spacer(1, 1 * cm), p("Брянск 2026", "Center"), PageBreak()]

    story += [p("ИНДИВИДУАЛЬНОЕ ЗАДАНИЕ", "TitleRu"), p("по учебной практике по технологии разработки программного обеспечения", "Center")]
    for line in ["Обучающийся: ________________________________", "Руководитель практики от вуза: ________________________________", "Сроки прохождения практики: ________________________________", "Место прохождения практики: кафедра ИТ БГИТУ"]:
        story.append(p(line))
    story.append(p("Содержание индивидуального задания:", "H2Ru"))
    story.append(table([
        ["Раздел", "Содержание"],
        ["Введение", "Актуальность, цель и задачи разработки приложения прогноза погоды."],
        ["Общая характеристика", "Назначение, функции, сценарии пользователя и схема работы."],
        ["Структура кода", "Фактическая структура weather_forecast_app и назначение файлов."],
        ["Серверная часть", "FastAPI, маршруты, Open-Meteo, обработка ошибок."],
        ["Клиентская часть", "HTML, JavaScript, localStorage, локализация, поиск и настройки."],
        ["Интерфейс", "CSS, premium minimal UI, glassmorphism, погодные фоны, адаптивность."],
        ["Тестирование", "Проверка запуска, API, статических файлов и пользовательских сценариев."],
    ], [4.5 * cm, 11.9 * cm]))
    story += [PageBreak(), p("СОДЕРЖАНИЕ", "TitleRu")]
    for i, title in enumerate(SECTIONS, 4):
        story.append(p(f"{title} ................................................................................ {i}"))
    story.append(PageBreak())

    story.append(p("ВВЕДЕНИЕ", "H1Ru"))
    for para in [
        "Целью учебной практики является получение практического опыта в проектировании, разработке и описании программного решения на примере веб-приложения прогноза погоды. В ходе работы был разработан проект Weather Forecast App, который получает данные из открытых погодных API, отображает текущую погоду и прогноз, поддерживает выбор города, локализацию интерфейса и обработку сетевых ошибок.",
        "Актуальность темы связана с тем, что погодные сервисы относятся к массовым информационным приложениям. Пользователь ожидает быстрый доступ к температуре, влажности, ветру, прогнозу на несколько дней, поиску города и понятному поведению приложения при отсутствии интернета.",
        "Объектом разработки является веб-приложение Weather Forecast App. Предметом разработки являются серверные маршруты FastAPI, интеграция с Open-Meteo, HTML-разметка интерфейса, CSS-визуальная система, клиентская логика JavaScript и сценарии тестирования.",
        "Цель проекта - разработать приложение для получения прогноза погоды по выбранному городу с современным мобильным интерфейсом, настройками пользователя и устойчивой обработкой ошибок сетевого взаимодействия.",
    ]:
        story.append(p(para))
    story.append(p("Для достижения цели были поставлены следующие задачи:"))
    story.append(bullets(GOALS))

    story.append(p("1. ОБЩАЯ ХАРАКТЕРИСТИКА ПРОЕКТА", "H1Ru"))
    story.append(p("Weather Forecast App - учебное веб-приложение прогноза погоды. Оно построено по схеме «браузер - FastAPI backend - Open-Meteo API». Пользователь выбирает город и получает текущую температуру, влажность, ветер, почасовой прогноз и прогноз на 10 дней."))
    story.append(p("Отчёт приведён в соответствие с фактической реализацией проекта: серверная часть находится в main.py, HTML-шаблон - в templates/index.html, клиентская логика - в static/app.js, оформление - в static/styles.css, а погодные изображения - в static/assets."))
    story.append(p("Основные функциональные возможности", "H2Ru"))
    story.append(table(FEATURE_ROWS, [5.1 * cm, 11.3 * cm]))
    story.append(p("Общая схема работы", "H2Ru"))
    story.append(p("Браузер -> templates/index.html -> static/styles.css и static/app.js -> /api/search или /api/weather -> FastAPI main.py -> Open-Meteo Geocoding API / Forecast API -> JSON-ответ -> отрисовка текущей погоды, почасового прогноза, прогноза на 10 дней и карточек городов.", "Code"))

    story.append(p("2. АНАЛИЗ СТРУКТУРЫ ИСХОДНОГО КОДА", "H1Ru"))
    story.append(p("Проект имеет компактную структуру, удобную для учебной практики и демонстрации. Серверный код, шаблон, статические файлы и локальные погодные изображения разделены по назначению."))
    story.append(p("weather_forecast_app/\n├── main.py\n├── requirements.txt\n├── README.md\n├── templates/\n│   └── index.html\n└── static/\n    ├── app.js\n    ├── styles.css\n    └── assets/\n        ├── weather-clear-day.png\n        ├── weather-clear-night.png\n        ├── weather-cloudy.png\n        ├── weather-rain-clouds.png\n        └── weather-rain-night.png", "CodeRu"))
    story.append(table(FILE_ROWS, [4.8 * cm, 11.6 * cm]))
    story.append(p("Технологический стек", "H2Ru"))
    story.append(table(TECH_ROWS, [5.0 * cm, 11.4 * cm]))

    story.append(p("3. РАЗРАБОТКА СЕРВЕРНОЙ ЧАСТИ", "H1Ru"))
    story.append(p("Backend реализован в main.py на FastAPI. Файл создаёт экземпляр приложения, подключает StaticFiles, подключает Jinja2Templates, содержит словарь WMO-кодов погоды, функции получения JSON, геокодирования и маршруты API."))
    story.append(p('app = FastAPI(title="Weather Forecast App", version="1.0.0")\napp.mount("/static", StaticFiles(directory="static"), name="static")\ntemplates = Jinja2Templates(directory="templates")', "Code"))
    story.append(table(API_ROWS))
    story.append(p("Функция geocode_city обращается к Open-Meteo Geocoding API. Она принимает название города и language. Если выбран русский язык, запрос «Москва» корректно возвращает русские результаты. Функция get_weather получает координаты города, затем обращается к Forecast API и формирует ответ с текущей погодой, 24 часами прогноза и прогнозом на 10 дней."))
    story.append(p("Для ошибок предусмотрены HTTPException: 404 при отсутствии города, 502 при недоступности внешнего сервиса и 504 при таймауте."))

    story.append(p("4. РАЗРАБОТКА КЛИЕНТСКОЙ ЧАСТИ", "H1Ru"))
    story.append(p("Клиентская часть реализована в templates/index.html и static/app.js. В одном HTML-шаблоне находятся четыре экрана: главная погода, поиск, настройки и ошибка сети. JavaScript переключает экраны, выполняет fetch-запросы, сохраняет настройки в localStorage и обновляет DOM."))
    story.append(p("Поиск города реализован с debounce-задержкой и поддерживает русский и английский язык. При выборе города приложение получает прогноз и возвращает пользователя на главный экран. В настройках можно переключать уведомления, язык, единицы измерения, тему и город по умолчанию."))
    story.append(p("Для безопасности строковые данные экранируются функцией escapeHtml. Погодные SVG-иконки формируются внутри приложения, а пользовательские названия городов перед выводом проходят экранирование."))

    story.append(p("5. ОФОРМЛЕНИЕ И АДАПТИВНЫЙ ИНТЕРФЕЙС", "H1Ru"))
    story.append(p("Интерфейс выполнен как современное мобильное погодное приложение: крупная температура, минималистичная типографика, стеклянные карточки, тонкие рамки, мягкие тени и premium-нижняя навигация."))
    story.append(p("Фоны главного экрана и карточек городов выбираются из локальных изображений static/assets. Алгоритм анализирует WMO-код погоды и время суток: дождь/гроза получают дождевой фон, облачность - облачный фон, ясная погода днём - светлый фон, ясная погода ночью - фон с луной."))
    story.append(p("CSS поддерживает тёмную и светлую тему. Для светлой темы отдельно настроены поле поиска, настройки, retry-кнопка и навигация, чтобы элементы не сливались с фоном."))

    story.append(p("6. ТЕСТИРОВАНИЕ И КОНТРОЛЬ КАЧЕСТВА", "H1Ru"))
    story.append(p("Тестирование выполнялось на локальном сервере http://127.0.0.1:8000. Проверялись backend, статические ресурсы, JavaScript-синтаксис и пользовательские сценарии."))
    story.append(table(TEST_ROWS, [6.0 * cm, 10.4 * cm]))
    story.append(p("Проверенные сценарии: открытие приложения, поиск города на русском и английском языке, выбор города, смена °C/°F, смена языка, выбор города по умолчанию, отображение экрана Network error и доступность локальных погодных фонов."))

    story.append(p("7. ИТОГИ И ВЫВОДЫ", "H1Ru"))
    story.append(p("В ходе учебной практики разработано веб-приложение Weather Forecast App, соответствующее требованиям задачи «Прогноз погоды». Приложение получает данные через API, отображает текущую температуру, влажность и ветер, показывает почасовой прогноз и прогноз на 10 дней, позволяет выбирать город и обрабатывает сетевые ошибки."))
    story.append(p("Проект демонстрирует полный цикл разработки программного модуля: проектирование требований, создание HTML/CSS/JavaScript-интерфейса, написание Python backend, интеграцию с внешним API и тестирование. Для дальнейшего развития рекомендуется добавить Pydantic-модели, кэширование, модульные тесты, избранные города и desktop-версию интерфейса."))

    story.append(p("8. ПРИЛОЖЕНИЕ", "H1Ru"))
    story.append(p('cd "C:\\Users\\siefi\\Documents\\New project\\weather_forecast_app"\npip install -r requirements.txt\nuvicorn main:app --reload\n# открыть http://127.0.0.1:8000', "CodeRu"))
    story.append(table([
        ["Требование", "Статус"],
        ["Получение данных о погоде через API", "Выполнено: Open-Meteo Forecast API."],
        ["Текущая температура, влажность, ветер", "Выполнено на главном экране."],
        ["Прогноз на несколько дней", "Выполнено: 10-дневный прогноз."],
        ["Выбор города", "Выполнено: экран поиска и Default City."],
        ["Ошибка сети", "Выполнено: экран Network error и Retry."],
        ["HTML, CSS, Python", "Выполнено: index.html, styles.css/app.js, main.py."],
    ], [6.0 * cm, 10.4 * cm]))

    doc.build(story)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Heading 1"].font.color.rgb = RGBColor(31, 77, 120)
    doc.styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)

    def para(text="", style=None, bold=False, center=False):
        p = doc.add_paragraph(style=style)
        r = p.add_run(text)
        r.bold = bold
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def tbl(rows):
        table = doc.add_table(rows=1, cols=len(rows[0]))
        table.style = "Table Grid"
        for i, val in enumerate(rows[0]):
            table.rows[0].cells[i].text = val
            table.rows[0].cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for row in rows[1:]:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = val
                cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for line in [
        "Министерство науки и высшего образования Российской Федерации",
        "ФГБОУ ВО «Брянский государственный инженерно-технологический университет»",
        "Многопрофильный колледж",
        "Кафедра «Информационные технологии»",
    ]:
        para(line, center=True)
    para("ОТЧЕТ", bold=True, center=True)
    para("по учебной практике по технологии разработки программного обеспечения", center=True)
    para("Тема проекта: «Weather Forecast App - веб-приложение прогноза погоды»", center=True)
    para(r"Локальный проект: C:\Users\siefi\Documents\New project\weather_forecast_app", center=True)
    for line in ["Студент ______________________________", "Группа ______________________________", "№ зачётной книжки ______________________________", "Руководитель от вуза ______________________________", "Нормоконтроль ______________________________"]:
        para(line)
    doc.add_page_break()

    para("СОДЕРЖАНИЕ", bold=True, center=True)
    for i, title in enumerate(SECTIONS, 4):
        para(f"{title} ................................................................................ {i}")
    doc.add_page_break()

    blocks = [
        ("ВВЕДЕНИЕ", [
            "Целью учебной практики является получение практического опыта в проектировании, разработке и описании программного решения на примере веб-приложения прогноза погоды.",
            "Разработан проект Weather Forecast App, который получает данные из Open-Meteo API, отображает текущую погоду и прогноз, поддерживает выбор города, локализацию интерфейса и обработку сетевых ошибок.",
        ]),
        ("1. ОБЩАЯ ХАРАКТЕРИСТИКА ПРОЕКТА", ["Weather Forecast App построен по схеме «браузер - FastAPI backend - Open-Meteo API» и реализует все требования практической задачи."]),
        ("2. АНАЛИЗ СТРУКТУРЫ ИСХОДНОГО КОДА", ["Фактическая структура проекта: main.py, templates/index.html, static/app.js, static/styles.css, static/assets, requirements.txt, README.md."]),
        ("3. РАЗРАБОТКА СЕРВЕРНОЙ ЧАСТИ", ["Backend реализован в main.py на FastAPI. Доступны маршруты /, /api/search и /api/weather."]),
        ("4. РАЗРАБОТКА КЛИЕНТСКОЙ ЧАСТИ", ["Клиентская часть реализована в index.html и app.js: переключение экранов, поиск, localStorage, темы, локализация, выбор города и отрисовка прогноза."]),
        ("5. ОФОРМЛЕНИЕ И АДАПТИВНЫЙ ИНТЕРФЕЙС", ["Интерфейс выполнен в мобильном premium minimal стиле со стеклянными карточками, нижней навигацией и погодными фонами."]),
        ("6. ТЕСТИРОВАНИЕ И КОНТРОЛЬ КАЧЕСТВА", ["Проверены запуск приложения, API, статика, русский поиск, JavaScript и основные сценарии пользователя."]),
        ("7. ИТОГИ И ВЫВОДЫ", ["Проект соответствует требованиям: API погоды, текущие показатели, прогноз, выбор города и ошибка сети реализованы."]),
        ("8. ПРИЛОЖЕНИЕ", ['Запуск: pip install -r requirements.txt; uvicorn main:app --reload; открыть http://127.0.0.1:8000.']),
    ]
    for title, paras in blocks:
        para(title, style="Heading 1")
        for item in paras:
            para(item)
        if title.startswith("1."):
            para("Основные функциональные возможности", style="Heading 2")
            tbl(FEATURE_ROWS)
        if title.startswith("2."):
            para("Назначение файлов", style="Heading 2")
            tbl(FILE_ROWS)
            para("Технологический стек", style="Heading 2")
            tbl(TECH_ROWS)
        if title.startswith("3."):
            para("API-маршруты", style="Heading 2")
            tbl(API_ROWS)
        if title.startswith("6."):
            tbl(TEST_ROWS)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_pdf()
    build_docx()
    print(PDF_PATH)
    print(DOCX_PATH)
